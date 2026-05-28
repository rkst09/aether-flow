from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def _label(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value or "—")


def generate_ba_docx(
    doc_data: dict,
    audit_data: dict,
    copy_suggestions: list,
    output_path: str,
):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title page ────────────────────────────────────────────────────────────
    title = doc.add_heading(doc_data.get("project_name", "Untitled Project"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("BA Handoff Document — Design Intelligence Output")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    _heading(doc, "Executive Summary", 1)
    doc.add_paragraph(doc_data.get("summary", ""))
    doc.add_paragraph()

    # ── Screen Documentation ──────────────────────────────────────────────────
    _heading(doc, "Screen Documentation", 1)

    screens = doc_data.get("screens", [])
    for screen in screens:
        doc.add_paragraph()
        _heading(doc, screen.get("name", "Unnamed Screen"), 2)

        _label(doc, "Type", screen.get("type", ""))
        _label(doc, "Primary Persona", screen.get("persona", ""))
        doc.add_paragraph()

        doc.add_paragraph(screen.get("annotation", ""))
        doc.add_paragraph()

        items = screen.get("interaction_logic", [])
        if items:
            _heading(doc, "Interaction Logic", 3)
            for item in items:
                _bullet(doc, item)

        items = screen.get("edge_cases", [])
        if items:
            _heading(doc, "Edge Cases", 3)
            for item in items:
                _bullet(doc, item)

        items = screen.get("validation_rules", [])
        if items:
            _heading(doc, "Validation Rules", 3)
            for item in items:
                _bullet(doc, item)

        items = screen.get("navigation_flows", [])
        if items:
            _heading(doc, "Navigation Flows", 3)
            for item in items:
                _bullet(doc, item)

    # ── UX Audit Summary ──────────────────────────────────────────────────────
    if audit_data:
        doc.add_page_break()
        _heading(doc, "UX Audit Findings", 1)

        sections_map = {
            "friction_points":    "Friction Points",
            "accessibility_issues": "Accessibility Issues",
            "missing_states":     "Missing States",
            "consistency_issues": "Consistency Issues",
        }

        for key, label in sections_map.items():
            items = audit_data.get(key, [])
            if not items:
                continue
            _heading(doc, label, 2)
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run(f"[{item.get('priority', 'medium').upper()}] {item.get('title', '')}")
                run.bold = True
                doc.add_paragraph(item.get("description", ""))
                screen_ref = item.get("screen_affected") or ", ".join(item.get("screens_affected", []))
                if screen_ref:
                    _label(doc, "Screen", screen_ref)
                doc.add_paragraph()

    # ── Copy Review Summary ───────────────────────────────────────────────────
    if copy_suggestions:
        doc.add_page_break()
        _heading(doc, "UX Copywriting Review", 1)

        current_screen = None
        for item in copy_suggestions:
            screen = item.get("screen", "")
            if screen != current_screen:
                current_screen = screen
                _heading(doc, screen, 2)

            p = doc.add_paragraph()
            run = p.add_run(f"[{item.get('category', '').replace('_', ' ').title()}]")
            run.bold = True
            run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

            _label(doc, "Before", item.get("original", ""))
            _label(doc, "After", item.get("suggested", ""))
            _label(doc, "Why", item.get("rationale", ""))
            doc.add_paragraph()

    doc.save(output_path)


def generate_persona_docs_docx(
    persona_docs: list[dict],
    output_path: str,
    project_name: str = "Aether Project",
):
    """Generate a real .docx handoff from the rich Phase 06 persona_docs payload."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading(project_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("BA Handoff Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    subtitle.runs[0].font.size = Pt(11)

    for persona in persona_docs:
        doc.add_page_break()
        _heading(doc, f"{persona.get('name', 'Persona')} — {persona.get('role', 'User')}", 1)
        _label(doc, "Initials", persona.get("initial", "—"))
        _label(doc, "Context", persona.get("context", ""))

        goals = persona.get("goals", []) or []
        if goals:
            _heading(doc, "Goals", 2)
            for goal in goals:
                _bullet(doc, goal)

        for module in persona.get("modules", []) or []:
            _heading(doc, module.get("name", "Module"), 2)
            doc.add_paragraph(module.get("purpose", ""))

            flows = module.get("flows", []) or []
            if flows:
                _heading(doc, "Flows", 3)
                for flow in flows:
                    _bullet(doc, flow)

            for screen in module.get("screens", []) or []:
                _heading(doc, screen.get("name", "Screen"), 3)
                _label(doc, "Screen Type", screen.get("screenType", ""))
                _label(doc, "Purpose", screen.get("purpose", ""))
                _label(doc, "Personas", ", ".join(screen.get("personas", []) or []) or "—")
                _label(doc, "Completeness", f"{screen.get('completeness', 0)}%")

                elements = screen.get("elements", []) or []
                if elements:
                    _heading(doc, "Elements", 4)
                    for element in elements:
                        conditions = element.get("conditions")
                        suffix = f" ({conditions})" if conditions else ""
                        _bullet(
                            doc,
                            f"{element.get('name', 'Element')} [{element.get('elType', 'Component')}] — {element.get('function', '')}{suffix}",
                        )

                interactions = screen.get("interactions", []) or []
                if interactions:
                    _heading(doc, "Interactions", 4)
                    for interaction in interactions:
                        nav = interaction.get("navigation")
                        nav_suffix = f" Navigation: {nav}." if nav else ""
                        _bullet(doc, f"{interaction.get('trigger', '')} → {interaction.get('result', '')}.{nav_suffix}")

                states = screen.get("states", []) or []
                if states:
                    _heading(doc, "States", 4)
                    for state in states:
                        _bullet(doc, f"{state.get('state', 'State')}: {state.get('description', '')}")

                validation = screen.get("validation", []) or []
                if validation:
                    _heading(doc, "Validation", 4)
                    for rule in validation:
                        required = "Required" if rule.get("required") else "Optional"
                        parts = [required]
                        if rule.get("format"):
                            parts.append(f"Format: {rule['format']}")
                        if rule.get("constraints"):
                            parts.append(f"Constraints: {rule['constraints']}")
                        if rule.get("errorBehavior"):
                            parts.append(f"Error: {rule['errorBehavior']}")
                        _bullet(doc, f"{rule.get('field', 'Field')} — {' | '.join(parts)}")

                navigation = screen.get("navigation", {}) or {}
                nav_lines = []
                if navigation.get("entry"):
                    nav_lines.append(f"Entry: {', '.join(navigation['entry'])}")
                if navigation.get("exit"):
                    nav_lines.append(f"Exit: {', '.join(navigation['exit'])}")
                if navigation.get("conditional"):
                    nav_lines.append(f"Conditional: {', '.join(navigation['conditional'])}")
                if nav_lines:
                    _heading(doc, "Navigation", 4)
                    for line in nav_lines:
                        _bullet(doc, line)

                edge_cases = screen.get("edgeCases", []) or []
                if edge_cases:
                    _heading(doc, "Edge Cases", 4)
                    for edge_case in edge_cases:
                        _bullet(doc, edge_case)

                copy_refs = screen.get("copyRefs", []) or []
                if copy_refs:
                    _heading(doc, "Copy References", 4)
                    for copy_ref in copy_refs:
                        _bullet(
                            doc,
                            f"{copy_ref.get('text', '')} [{copy_ref.get('kind', 'static')}] — {copy_ref.get('purpose', '')}",
                        )

                dependencies = screen.get("dependencies", []) or []
                if dependencies:
                    _heading(doc, "Dependencies", 4)
                    for dependency in dependencies:
                        _bullet(doc, dependency)

                flags = screen.get("flags", []) or []
                if flags:
                    _heading(doc, "Flags", 4)
                    for flag in flags:
                        _bullet(doc, flag)

    doc.save(output_path)
